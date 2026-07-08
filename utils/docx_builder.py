"""
Builds the PMI Closeout Report by modifying the original template directly.
Preserves all original formatting, fonts, tables, branding, header/footer.
"""

import io
import os
import json
import re
import zipfile
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = Path(__file__).parent.parent / "TN20133-A PMI Closeout Report .docx"

TOWER_TYPE_DISPLAY = {
    "Self Support": "Self Support Tower",
    "Guyed":        "Guyed Tower",
    "Monopole":     "Monopole",
}

# Field observation section names per tower type (exported for use in app.py)
OBS_SECTIONS_BY_TOWER = {
    "Self Support": [
        "Structural Member Verification",
        "Connection and Installation Verification",
        "Modification Installation",
        "Connection and Welding",
        "Alignment and Eccentricity",
        "Coating and Protection",
        "Interference Check",
        "Final Verification",
    ],
    "Guyed": [
        "Structural Member Verification",
        "Connection and Installation Verification",
        "Modification Installation",
        "Alignment and Eccentricity",
        "Interference Check",
        "Guy Wire Verification",
        "Final Verification",
    ],
    "Monopole": [
        "Structural Member Verification",
        "Connection and Installation Verification",
        "Modification Installation",
        "Connection and Welding",
        "Alignment and Eccentricity",
        "Coating and Protection",
        "Interference Check",
        "Final Verification",
    ],
}

# Photo Documentation table — special section labels per tower type
_PHOTO_TABLE_SPECIALS = {
    "Self Support": [
        "Overall Tower Views – Completed Modifications",
        "Field Measurements and Verification of Installed Structural Members",
    ],
    "Guyed": [
        "Tension Gauge Photographs – Guy Wire Anchor Points",
        "Overall Tower Views – Completed Modifications",
    ],
    "Monopole": [
        "Overall Tower Views – Completed Modifications",
        "Field Measurements and Verification of Installed Structural Members",
    ],
}


# ─── SMART TEXT REPLACEMENT ──────────────────────────────────────────────────

def _smart_replace_in_para(para, old: str, new: str) -> bool:
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False

    start = full.index(old)
    end   = start + len(old)
    pos = 0
    start_run_idx = start_off = None
    end_run_idx   = end_off   = None

    for i, run in enumerate(para.runs):
        run_end = pos + len(run.text)
        if start_run_idx is None and start < run_end:
            start_run_idx, start_off = i, start - pos
        if end_run_idx is None and end <= run_end:
            end_run_idx, end_off = i, end - pos
            break
        pos = run_end

    if start_run_idx is None or end_run_idx is None:
        return False

    runs = para.runs
    if start_run_idx == end_run_idx:
        r = runs[start_run_idx]
        r.text = r.text[:start_off] + new + r.text[end_off:]
    else:
        runs[start_run_idx].text = runs[start_run_idx].text[:start_off] + new
        for i in range(start_run_idx + 1, end_run_idx):
            runs[i].text = ""
        runs[end_run_idx].text = runs[end_run_idx].text[end_off:]

    return True


def _apply_replacements(doc, replacements: dict):
    for para in doc.paragraphs:
        for old, new in replacements.items():
            _smart_replace_in_para(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        _smart_replace_in_para(para, old, new)


# ─── TABLE HELPERS ───────────────────────────────────────────────────────────

def _find_table_by_header(doc, *header_texts):
    for table in doc.tables:
        row0 = " ".join(c.text for c in table.rows[0].cells)
        if all(h in row0 for h in header_texts):
            return table
    return None


def _set_cell_text(cell, text: str):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _rebuild_table(doc, header_texts: tuple, new_rows: list):
    table = _find_table_by_header(doc, *header_texts)
    if table is None:
        return
    template_tr = deepcopy(table.rows[1]._tr) if len(table.rows) > 1 else None
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)
    for row_values in new_rows:
        if template_tr is not None:
            new_tr = deepcopy(template_tr)
            table._tbl.append(new_tr)
            new_row = table.rows[-1]
            for i, val in enumerate(row_values):
                if i < len(new_row.cells):
                    _set_cell_text(new_row.cells[i], str(val))
        else:
            new_row_cells = table.add_row().cells
            for i, val in enumerate(row_values):
                if i < len(new_row_cells):
                    new_row_cells[i].text = str(val)


# ─── COVER PAGE ADDRESS ───────────────────────────────────────────────────────

def _replace_para_value_after_label(para, label: str, new_value: str) -> bool:
    full = "".join(r.text for r in para.runs)
    if label not in full:
        return False
    colon_end = full.index(label) + len(label)
    spaces_end = colon_end
    while spaces_end < len(full) and full[spaces_end] == " ":
        spaces_end += 1
    label_with_spaces = full[:spaces_end]
    new_full = label_with_spaces + new_value
    runs = para.runs
    if runs:
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""
    return True


def _replace_cover_address(doc, site_address: str, gps: str):
    paras = list(doc.paragraphs)
    addr_idx = None
    for i, para in enumerate(paras):
        if "Site Address:" in para.text:
            addr_idx = i
            break
    if addr_idx is None:
        return

    address_lines = [l for l in site_address.split("\n") if l.strip()]
    if len(address_lines) == 1 and "," in address_lines[0]:
        # No explicit line break typed (e.g. AI-extracted address arrives as
        # one string) — split on the first comma so street address and
        # city/state/zip still land on their own template lines.
        street, rest = address_lines[0].split(",", 1)
        address_lines = [street.strip(), rest.strip(" ,")]
    _replace_para_value_after_label(paras[addr_idx], "Site Address:", address_lines[0] if address_lines else "")

    for offset, line in enumerate([
        address_lines[1] if len(address_lines) > 1 else "",
        gps or ""
    ], start=1):
        if addr_idx + offset < len(paras):
            p = paras[addr_idx + offset]
            full = "".join(r.text for r in p.runs)
            if full.startswith("   "):
                stripped = full.lstrip(" ")
                indent = full[: len(full) - len(stripped)]
                new_full = indent + line
                if p.runs:
                    p.runs[0].text = new_full
                    for r in p.runs[1:]:
                        r.text = ""


# ─── SCOPE BULLETS ───────────────────────────────────────────────────────────

def _replace_scope_bullets(doc, modifications: list):
    body = doc.element.body
    body_paras = [c for c in body if c.tag == qn("w:p")]
    start_idx = end_idx = None
    template_bullet = None

    for i, p in enumerate(body_paras):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if "consisted of the following items" in text:
            start_idx = i
        elif start_idx is not None and template_bullet is None:
            template_bullet = deepcopy(p)
        elif "Based on our field observation" in text and start_idx is not None:
            end_idx = i
            break

    if start_idx is None or end_idx is None or template_bullet is None:
        return

    for p in body_paras[start_idx + 1: end_idx]:
        p.getparent().remove(p)

    anchor = body_paras[start_idx]
    for mod in reversed(modifications):
        new_p = deepcopy(template_bullet)
        for t in new_p.iter(qn("w:t")):
            t.text = ""
        t_list = list(new_p.iter(qn("w:t")))
        if t_list:
            t_list[0].text = mod["description"]
        anchor.addnext(new_p)


# ─── DEFICIENCIES ─────────────────────────────────────────────────────────────

def _set_deficiencies(doc, deficiencies_text: str, no_deficiencies: bool):
    if no_deficiencies or not deficiencies_text.strip():
        return
    body = doc.element.body
    for p in body:
        if p.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if "No visible deficiencies were observed" in text:
            for t in p.iter(qn("w:t")):
                t.text = ""
            t_list = list(p.iter(qn("w:t")))
            if t_list:
                t_list[0].text = deficiencies_text
            break


# ─── FIELD OBSERVATIONS ──────────────────────────────────────────────────────

def generate_field_observations(modifications: list, tower_type: str = "Self Support") -> dict:
    """
    Call Gemini to generate 2 observation bullets per section.
    Exported for use in app.py (Step 5 preview).

    Raises on API failure (quota/billing/network/etc.) instead of swallowing
    the error, so callers — especially the interactive "Generate AI
    Observations" button in app.py — can show the user what actually went
    wrong instead of silently rendering blank text areas.
    """
    api_key = os.environ.get("GEMINI_API_KEY", "")
    if not api_key or not modifications:
        return {}

    sections = OBS_SECTIONS_BY_TOWER.get(tower_type, OBS_SECTIONS_BY_TOWER["Self Support"])
    mod_text = "\n".join(
        f"- {m['mod_id']}: {m['description']} at elevation {m['elevation']}"
        for m in modifications
    )

    sections_template = "{" + ",\n  ".join(f'"{s}": ["point 1", "point 2"]' for s in sections) + "}"

    if tower_type == "Guyed":
        context = (
            "This is a guyed tower PMI closeout and inspection report. "
            "The tower uses guy wires, anchors, and a central mast. "
            "Reference guy system hardware, anchor conditions, and foundations where relevant."
        )
    else:
        context = "This is a self-support lattice tower PMI closeout report."

    prompt = f"""You are writing the Field Observation Details section for a structural engineering PMI closeout report.

{context}

Modifications performed on this tower:
{mod_text}

For each section below, write EXACTLY 2 concise bullet points. Use passive voice, professional structural engineering language. Reference actual member types, elevations, or modification IDs where appropriate. Do NOT write generic statements.

Return ONLY valid JSON (no markdown fences, no explanation):
{sections_template}"""

    from google import genai
    from agents.retry import generate_content_with_retry
    client = genai.Client(api_key=api_key)
    response = generate_content_with_retry(client, model="gemini-2.5-flash", contents=prompt)
    text = response.text.strip()
    text = re.sub(r"^```json\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def _set_para_text(p, text: str):
    for t in p.iter(qn("w:t")):
        t.text = ""
    t_list = list(p.iter(qn("w:t")))
    if t_list:
        t_list[0].text = text


def _replace_field_observations(doc, observations: dict):
    if not observations:
        return

    body = doc.element.body
    body_paras = [c for c in body if c.tag == qn("w:p")]

    fod_idx = obs_def_idx = None
    heading_tmpl = bullet_tmpl = None

    for i, p in enumerate(body_paras):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if "Field Observation Details" in text and fod_idx is None:
            fod_idx = i
        elif fod_idx is not None:
            if "Structural Member Verification" in text and heading_tmpl is None:
                heading_tmpl = deepcopy(p)
            elif heading_tmpl is not None and bullet_tmpl is None:
                stripped = text.strip()
                if stripped:
                    bullet_tmpl = deepcopy(p)
            # Exact-match the heading paragraph itself — a loose substring
            # check on "Observed Deficiencies" also matches the template's
            # own "Final Verification" boilerplate bullet ("...except as
            # noted in the Observed Deficiencies section."), which stops the
            # removal range one paragraph early and leaks that bullet through
            # alongside the AI-generated ones.
            if text.strip() in ("Observed Deficiencies", "Observed Deficiencies:"):
                obs_def_idx = i
                break

    if fod_idx is None or obs_def_idx is None or heading_tmpl is None:
        return

    if bullet_tmpl is None:
        bullet_tmpl = deepcopy(heading_tmpl)

    fod_para = body_paras[fod_idx]

    for p in body_paras[fod_idx + 1: obs_def_idx]:
        p.getparent().remove(p)

    all_paras = []

    intro = deepcopy(bullet_tmpl)
    _set_para_text(intro, "The following observations were made during the site visit in accordance with the approved modification drawings.")
    all_paras.append(intro)

    for section, bullets in observations.items():
        hp = deepcopy(heading_tmpl)
        _set_para_text(hp, section + ":")
        all_paras.append(hp)
        for bullet_text in bullets:
            bp = deepcopy(bullet_tmpl)
            _set_para_text(bp, bullet_text)
            all_paras.append(bp)

    for p in reversed(all_paras):
        fod_para.addnext(p)


# ─── TRUNCATE ─────────────────────────────────────────────────────────────────

def _truncate_after(doc, marker_text: str):
    body = doc.element.body
    found = False
    to_remove = []

    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if found:
            to_remove.append(child)
        elif child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if text in (marker_text, marker_text + ":"):
                found = True

    for elem in to_remove:
        body.remove(elem)


# ─── HEADING STYLE TEMPLATE (for appended sections) ──────────────────────────

def _get_heading_pPr_template(doc, marker_text: str):
    """
    Deep-copies the <w:pPr> (style + auto-number list) of a real "Heading 1"
    paragraph in the template — e.g. "On-Site Inspection Photographs:" — so
    sections we append later (extra documents, custom certificates) can reuse
    the exact same numbered-heading formatting and continue the template's
    Word list numbering. Without this, appended sections render as plain bold
    text instead of real headings, and the document's Table of Contents field
    (which tracks Heading-1..3 styles) never picks them up.
    """
    for p in doc.paragraphs:
        if p.text.strip() in (marker_text, marker_text + ":") and p.style.name == "Heading 1":
            pPr = p._p.find(qn("w:pPr"))
            return deepcopy(pPr) if pPr is not None else None
    return None


# ─── BLANK PARAGRAPH CLEANUP ──────────────────────────────────────────────────

def _collapse_excess_blank_paragraphs(doc, max_blanks: int = 1):
    """
    The master template has hand-typed blank paragraphs (a handful up to 16 in
    a row) used as manual spacing before some section headings, sized for the
    template's original (long) sample content. When our dynamically generated
    content is shorter, those blank paragraphs land stacked at the top of a
    page instead of acting as a small visual gap, producing a near-empty page
    before headings like "Observed Deficiencies:". This trims any run of more
    than `max_blanks` consecutive empty paragraphs immediately before a
    Heading 1/2 paragraph, anywhere in the document.
    """
    paras = doc.paragraphs

    for i, p in enumerate(paras):
        if p.style.name not in ("Heading 1", "Heading 2") or not p.text.strip():
            continue
        # Walk backward over ANY blank paragraph — including empty leftover
        # "Heading 1" placeholders in the template — stopping at real content.
        run = []
        j = i - 1
        while j >= 0 and paras[j].text.strip() == "":
            run.append(paras[j])
            j -= 1
        if len(run) > max_blanks:
            for blank_p in run[max_blanks:]:
                el = blank_p._p
                if el.getparent() is not None:
                    el.getparent().remove(el)


# ─── PDF TO IMAGES ────────────────────────────────────────────────────────────

def _pdf_to_images(pdf_bytes: bytes) -> list:
    try:
        import fitz
        pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        images = []
        for page in pdf_doc:
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
            images.append(pix.tobytes("png"))
        pdf_doc.close()
        return images
    except Exception:
        return []


# ─── PHOTO SIZE CONSTRAINT ────────────────────────────────────────────────────

def _ensure_portrait(img_bytes: bytes) -> bytes:
    """
    Normalizes camera EXIF orientation, then rotates any still-landscape
    photo 90° so every photo in the report displays portrait — customer
    photo pools mix landscape and portrait shots, which looked inconsistent
    in the photo grid.
    """
    try:
        from PIL import Image as PILImage, ImageOps
        img = PILImage.open(io.BytesIO(img_bytes))
        fmt = (img.format or "JPEG").upper()
        if fmt not in ("JPEG", "PNG"):
            fmt = "JPEG"
        img = ImageOps.exif_transpose(img)
        if img.width > img.height:
            img = img.rotate(-90, expand=True)
        if fmt == "JPEG" and img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format=fmt)
        return buf.getvalue()
    except Exception:
        return img_bytes


def _constrained_picture_dims(img_bytes: bytes, max_w: float = 2.9, max_h: float = 2.6) -> tuple:
    """
    Returns (width_inches, height_inches) that fits the image within max_w × max_h
    while preserving aspect ratio. Constrains portrait images by height so that
    two rows of photos always fit on a standard page.
    """
    try:
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(img_bytes))
        w, h = img.size
        if w == 0 or h == 0:
            return max_w, max_h
        aspect = w / h
        box_aspect = max_w / max_h
        if aspect >= box_aspect:
            return max_w, round(max_w / aspect, 4)
        else:
            return round(aspect * max_h, 4), max_h
    except Exception:
        return max_w, max_h


# ─── ADD CONTENT HELPERS ──────────────────────────────────────────────────────

def _add_section_heading(doc, text: str, heading_pPr=None):
    """
    Adds a section heading. If `heading_pPr` (from _get_heading_pPr_template)
    is supplied, the paragraph gets the template's real numbered "Heading 1"
    formatting instead of plain bold text — so it matches the template's other
    section headings and is picked up by the document's Table of Contents
    field when Word refreshes it.
    """
    para = doc.add_paragraph()
    if heading_pPr is not None:
        new_pPr = deepcopy(heading_pPr)
        existing_pPr = para._p.find(qn("w:pPr"))
        if existing_pPr is not None:
            para._p.remove(existing_pPr)
        para._p.insert(0, new_pPr)
        para.add_run(text)
    else:
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
    return para


def _add_photo_group_heading(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return para


def _clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_photo_table(doc, photos: list):
    """
    Add photos in a 2-column borderless table, 2 per row (4 per page).
    Every photo is rotated to portrait (see _ensure_portrait) and then
    size-constrained to fit within a 2.9" × 2.6" box, guaranteeing two rows
    per page.
    """
    if not photos:
        return

    n_rows = (len(photos) + 1) // 2
    table = doc.add_table(rows=n_rows, cols=2)

    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    for i, (img_bytes, caption, _mime, *_rest) in enumerate(photos):
        img_bytes = _ensure_portrait(img_bytes)
        row_idx = i // 2
        col_idx = i % 2
        cell = table.rows[row_idx].cells[col_idx]

        img_para = cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            w_in, h_in = _constrained_picture_dims(img_bytes)
            img_para.add_run().add_picture(
                io.BytesIO(img_bytes),
                width=Inches(w_in),
                height=Inches(h_in),
            )
        except Exception:
            img_para.add_run("[Image could not be inserted]")

        cap_para = cell.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap_para.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

    if len(photos) % 2 == 1:
        last_row = table.rows[-1]
        empty_cell = last_row.cells[1]
        for p in empty_cell.paragraphs:
            for run in p.runs:
                run.text = ""

    doc.add_paragraph()


def _add_file(doc, file_bytes: bytes, filename: str, mime_type: str):
    if mime_type in ("image/jpeg", "image/png"):
        try:
            run = doc.add_paragraph().add_run()
            run.add_picture(io.BytesIO(file_bytes), width=Inches(6.0))
        except Exception:
            doc.add_paragraph(f"[Could not insert: {filename}]")
        doc.add_paragraph()

    elif mime_type == "application/pdf":
        images = _pdf_to_images(file_bytes)
        if images:
            for img_bytes in images:
                try:
                    run = doc.add_paragraph().add_run()
                    run.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
                except Exception:
                    doc.add_paragraph(f"[PDF page could not be inserted: {filename}]")
                doc.add_paragraph()
        else:
            doc.add_paragraph(f"[PDF attached: {filename} — could not render]")
            doc.add_paragraph()
    else:
        p = doc.add_paragraph(f"[Attached: {filename}]")
        p.runs[0].italic = True
        doc.add_paragraph()


def _get_positions(tower_type: str, num_guys: int = 3) -> list:
    if tower_type == "Self Support":
        return ["Leg A", "Leg B", "Leg C"]
    elif tower_type == "Guyed":
        return [f"Guy {i+1}" for i in range(num_guys)]
    return []


# ─── ORPHANED IMAGE CLEANUP ──────────────────────────────────────────────────

def _strip_orphaned_images(docx_bytes: bytes) -> bytes:
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
            doc_xml = zin.read("word/document.xml").decode("utf-8")
            referenced_rids = set(re.findall(r'r:embed="(rId\d+)"', doc_xml))

            rels_xml = zin.read("word/_rels/document.xml.rels").decode("utf-8")
            img_type = "relationships/image"
            all_img_rels = {}
            for m in re.finditer(
                r'<Relationship[^>]+Id="(rId\d+)"[^>]+Type="[^"]*' + img_type + r'"[^>]+Target="([^"]+)"',
                rels_xml,
            ):
                all_img_rels[m.group(1)] = m.group(2)

            orphan_targets = {v for k, v in all_img_rels.items() if k not in referenced_rids}
            orphan_rids    = {k for k, v in all_img_rels.items() if k not in referenced_rids}

            if not orphan_targets:
                return docx_bytes

            orphan_paths = {f"word/{t}" for t in orphan_targets}

            out = io.BytesIO()
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename in orphan_paths:
                        continue
                    data = zin.read(item.filename)
                    if item.filename == "word/_rels/document.xml.rels":
                        for rid in orphan_rids:
                            data = re.sub(
                                rb'<Relationship[^>]*Id="' + rid.encode() + rb'"[^>]*/?>',
                                b"",
                                data,
                            )
                    zout.writestr(item, data)

            return out.getvalue()
    except Exception:
        return docx_bytes


# ─── TABLE OF CONTENTS SYNC ───────────────────────────────────────────────────
#
# The template's TOC is a real Word field (`TOC \o "1-3" \h \z \u`) inside a
# content-control (w:sdt). `updateFields=true` (set at the end of
# build_report) makes desktop Word recompute it — with correct page numbers
# — the instant the file is opened. But that recompute only happens in real
# Word; anything that doesn't execute field updates (Google Docs import,
# browser/quick-look previews, some PDF converters) just displays the
# *cached* entries that were baked in when the master template was authored,
# which have nothing to do with this report's actual generated sections.
#
# This rewrites those cached entries in place, before the file is handed
# back for download, so the titles are correct everywhere — even before any
# field update runs. Page numbers are left for Word to fill in live (we have
# no layout engine here to compute real pagination), so entries are shown
# without a fabricated number rather than a plausible-looking wrong one.

_TOC_STYLE_BY_LEVEL = {1: "TOC1", 2: "TOC2", 3: "TOC2"}


def _collect_headings_for_toc(doc) -> list:
    """
    Returns [(level, text), ...] for every Heading 1/2/3 paragraph in the
    final document, in document order. Mirrors what the TOC field's own
    `\\o "1-3"` switch would pick up on a real Word refresh.
    """
    headings = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style is not None else ""
        if style_name in ("Heading 1", "Heading 2", "Heading 3"):
            text = p.text.strip()
            if text:
                headings.append((int(style_name[-1]), text))
    return headings


def _rewrite_toc_entries(doc, headings: list):
    """
    Replaces the TOC field's cached entry paragraphs (style TOC1/TOC2/TOC3,
    immediately following the "Table of Contents" TOCHeading paragraph) with
    one paragraph per heading in `headings`. The outer field's begin /
    instrText / separate runs (which live at the start of the *first* entry
    paragraph) and its closing paragraph (which holds only the field's `end`
    fldChar) are preserved untouched, so the field stays fully valid and
    still self-corrects in real Word.

    No-ops safely if the template's TOC structure doesn't match what's
    expected here, rather than risking a corrupted document.
    """
    if not headings:
        return

    body = doc.element.body
    # The TOC lives inside a content control (w:sdt > w:sdtContent), not as
    # a direct child of the body, so this must walk all descendants — a
    # plain findall(qn("w:p")) (direct children only) would silently miss
    # it entirely.
    all_paragraphs = list(body.iter(qn("w:p")))

    def _style_of(p):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            return None
        style = pPr.find(qn("w:pStyle"))
        return style.get(qn("w:val")) if style is not None else None

    def _has_fldchar_end(p):
        return any(
            fld.get(qn("w:fldCharType")) == "end"
            for fld in p.iter(qn("w:fldChar"))
        )

    toc_head_idx = None
    for i, p in enumerate(all_paragraphs):
        if _style_of(p) == "TOCHeading":
            toc_head_idx = i
            break
    if toc_head_idx is None:
        return  # template has no TOC — nothing to sync

    entry_start = toc_head_idx + 1
    entry_end = entry_start
    while (entry_end < len(all_paragraphs)
           and _style_of(all_paragraphs[entry_end]) in ("TOC1", "TOC2", "TOC3")):
        entry_end += 1

    if entry_end >= len(all_paragraphs) or not _has_fldchar_end(all_paragraphs[entry_end]):
        return  # structure isn't what we expect — leave it alone

    old_entries = all_paragraphs[entry_start:entry_end]
    if not old_entries:
        return

    # The outer field's begin/instrText/separate runs are the leading
    # non-pPr, non-hyperlink children of the first entry paragraph.
    first_entry = old_entries[0]
    field_open_runs = []
    for child in list(first_entry):
        if child.tag == qn("w:pPr"):
            continue
        if child.tag == qn("w:hyperlink"):
            break
        field_open_runs.append(child)

    def _build_entry_paragraph(level: int, text: str, include_field_open: bool):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        style_el = OxmlElement("w:pStyle")
        style_el.set(qn("w:val"), _TOC_STYLE_BY_LEVEL.get(level, "TOC1"))
        pPr.append(style_el)
        p.append(pPr)
        if include_field_open:
            for r in field_open_runs:
                p.append(deepcopy(r))
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        p.append(run)
        return p

    new_paragraphs = [
        _build_entry_paragraph(level, text, include_field_open=(i == 0))
        for i, (level, text) in enumerate(headings)
    ]

    insert_point = old_entries[0]
    for new_p in new_paragraphs:
        insert_point.addprevious(new_p)
    for old_p in old_entries:
        old_p.getparent().remove(old_p)


# Local names of the CT_Settings elements that the OOXML schema requires to
# come AFTER w:updateFields. Word's settings.xml parser is order-sensitive —
# inserting updateFields at position 0 (before e.g. w:zoom) is invalid and
# can make Word treat the part as needing repair, silently dropping the
# setting so the TOC never actually auto-refreshes on open.
_SETTINGS_AFTER_UPDATE_FIELDS = (
    "hdrShapeDefaults", "footnotePr", "endnotePr", "compat", "rsids", "mathPr",
    "themeFontLang", "clrSchemeMapping", "doNotAutoCompressPictures",
    "shapeDefaults", "decimalSymbol", "listSeparator", "docId",
    "defaultImageDpi", "discardImageEditingData", "conflictMode",
    "chartTrackingRefBased",
)


def _insert_update_fields(settings_el):
    """Sets updateFields=true at a schema-valid position in settings.xml (see above)."""
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    for child in settings_el:
        if child.tag.split("}")[-1] in _SETTINGS_AFTER_UPDATE_FIELDS:
            child.addprevious(update_fields)
            return
    settings_el.append(update_fields)


# ─── MAIN BUILD FUNCTION ─────────────────────────────────────────────────────

def build_report(data: dict) -> bytes:
    doc = Document(TEMPLATE_PATH)

    # Captured before any mutation so appended sections (extra documents,
    # custom certificates) can reuse the template's real numbered Heading 1
    # formatting instead of plain bold text.
    heading_pPr = _get_heading_pPr_template(doc, "On-Site Inspection Photographs")

    info           = data["info"]
    modifications  = data.get("modifications", [])
    tower_type     = data.get("tower_type", "Self Support")
    num_guys       = data.get("num_guys", 3)
    photos         = data.get("photos", {})
    special_photos = data.get("special_photos", {})
    extra_photos   = data.get("extra_photos", [])
    documents      = data.get("documents", {})
    extra_documents = data.get("extra_documents", [])
    deficiencies   = data.get("deficiencies", "")
    no_deficiencies = data.get("no_deficiencies", True)

    tower_display = TOWER_TYPE_DISPLAY.get(tower_type, tower_type)

    # ── 1. Text replacements ─────────────────────────────────────────────────
    replacements = {
        "June 9, 2026":                  info.get("report_date", ""),
        "SBA Communication Corporation": info.get("client", ""),
        "Hillvale, TN Tower":            info.get("site_name", ""),
        "AT&T":                          info.get("carrier_name", ""),
        "Self Support Tower":            tower_display,
        "300’":                     info.get("tower_height", ""),
        "300'":                          info.get("tower_height", ""),
        "08/06/2026":                    info.get("observation_date", ""),
        "01/08/2026":                    info.get("drawing_date", ""),
        "S-02, S-03":                    info.get("drawing_sheets", ""),
        "Vinco, Inc":                    info.get("general_contractor", ""),
    }

    client_first = info.get("client", "").split()[0] if info.get("client") else "Client"
    replacements["Dear SBA Team"] = f"Dear {client_first} Team"

    old_desc = (
        "Post modification structural reinforcement and verification of "
        "tower components as per approved design drawings."
    )
    new_desc = info.get("project_description", old_desc)
    if new_desc != old_desc:
        replacements[old_desc] = new_desc

    _apply_replacements(doc, replacements)
    _apply_replacements(doc, {"TN20133-A": info.get("site_number", "")})

    # ── 2. Cover page address ────────────────────────────────────────────────
    _replace_cover_address(doc, info.get("site_address", ""), info.get("gps_coords", ""))

    # ── 3. Scope bullets ─────────────────────────────────────────────────────
    _replace_scope_bullets(doc, modifications)

    # ── 4. Modification Summary table ────────────────────────────────────────
    _rebuild_table(
        doc,
        ("Modification ID", "Description", "Elevation"),
        [[m["mod_id"], m["description"], m["elevation"]] for m in modifications],
    )

    # ── 5. Design Documentation table ────────────────────────────────────────
    site_no   = info.get("site_number", "")
    draw_date = info.get("drawing_date", "")
    _rebuild_table(
        doc,
        ("Document(s)", "Remarks", "Source"),
        [["Tower Modification Drawings",
          f"ATSS, Project#:{site_no} Dated {draw_date}",
          "Advanced Tower Structural Solutions"]],
    )

    # ── 6. Photo Documentation table ─────────────────────────────────────────
    positions = _get_positions(tower_type, num_guys)
    photo_rows = []
    counter = 1
    for mod in modifications:
        mid   = mod["mod_id"]
        elev  = mod["elevation"]
        short = mod["description"][:55] + ("..." if len(mod["description"]) > 55 else "")
        if positions:
            for pos in positions:
                photo_rows.append([str(counter), f"{mid} ({elev}) – {pos} – {short}", ""])
                counter += 1
        else:
            photo_rows.append([str(counter), f"{mid} ({elev}) – {short}", ""])
            counter += 1

    for label in _PHOTO_TABLE_SPECIALS.get(tower_type, _PHOTO_TABLE_SPECIALS["Self Support"]):
        photo_rows.append([str(counter), label, ""])
        counter += 1

    for entry_name, photo_list in extra_photos:
        if entry_name and photo_list:
            photo_rows.append([str(counter), entry_name, ""])
            counter += 1

    _rebuild_table(doc, ("Photo No.", "Description"), photo_rows)

    # ── 7. Deficiencies ──────────────────────────────────────────────────────
    _set_deficiencies(doc, deficiencies, no_deficiencies)

    # ── 8. Field observations ────────────────────────────────────────────────
    field_obs = data.get("field_observations") or {}
    if not field_obs:
        try:
            field_obs = generate_field_observations(modifications, tower_type)
        except Exception:
            field_obs = {}
    if field_obs:
        _replace_field_observations(doc, field_obs)

    # Trim the template's oversized manual spacer paragraphs now that section
    # content has its real (usually shorter) length, so headings like
    # "Observed Deficiencies:" don't land under a near-empty page.
    _collapse_excess_blank_paragraphs(doc)

    # ── 9. Truncate after photo heading ──────────────────────────────────────
    _truncate_after(doc, "On-Site Inspection Photographs")

    # ── 10 & 11. Photo sections — each group starts on its own page so a
    # group's photos are never split 2-and-2 across a page boundary by
    # Word's natural pagination.
    first_group = True

    def _start_group():
        nonlocal first_group
        if not first_group:
            doc.add_page_break()
        first_group = False

    for mod in modifications:
        mid  = mod["mod_id"]
        elev = mod["elevation"]
        if positions:
            for pos in positions:
                photo_list = photos.get((mid, pos), [])
                if photo_list:
                    _start_group()
                    _add_photo_group_heading(doc, f"{mid} ({elev}) – {pos}:")
                    _add_photo_table(doc, photo_list)
        else:
            photo_list = photos.get((mid, ""), [])
            if photo_list:
                _start_group()
                _add_photo_group_heading(doc, f"{mid} ({elev}):")
                _add_photo_table(doc, photo_list)

    for label, photo_list in special_photos.items():
        if photo_list:
            _start_group()
            _add_photo_group_heading(doc, f"{label}:")
            _add_photo_table(doc, photo_list)

    for entry_name, photo_list in extra_photos:
        if entry_name and photo_list:
            _start_group()
            _add_photo_group_heading(doc, f"{entry_name}:")
            _add_photo_table(doc, photo_list)

    # ── 12. Supporting documents — tower-type-specific fixed checklist ───────
    if tower_type == "Guyed":
        doc_sections = [
            ("as_built_drawings",     "As-Built Drawings (EOR)"),
            ("material_cert",         "Material Certification Report"),
            ("packing_slips",         "Packing Slips"),
            ("tension_report",        "Tension Report"),
            ("plumb_twist_report",    "Plumb & Twist Report"),
            ("fabrication_submittal", "Fabrication Submittal Package"),
            ("cold_galv_letter",      "Cold Galvanization Letter"),
        ]
    else:
        doc_sections = [
            ("as_built_drawings",     "As-Built Drawings (EOR)"),
            ("material_cert",         "Material Certification Report"),
            ("fabrication_submittal", "Fabrication Submittal Package"),
            ("fabrication_letter",    "Fabrication Letter"),
            ("cold_galv_letter",      "Cold Galvanization Letter"),
        ]

    for doc_key, section_title in doc_sections:
        doc.add_page_break()
        _add_section_heading(doc, f"{section_title}:", heading_pPr)
        file_list = documents.get(doc_key, [])
        if file_list:
            for fb, fname, mime in file_list:
                _add_file(doc, fb, fname, mime)
        else:
            doc.add_paragraph("[Not provided]")

    # ── 13. Certificates & extra documents — user-named, not a fixed generic
    # "Certificates" bucket. Each entry is its own titled section (real
    # Heading 1, continuing the template's numbering), so it shows up
    # correctly both in the document body and in Word's Table of Contents
    # field once refreshed.
    for entry_name, file_list in extra_documents:
        if not entry_name or not file_list:
            continue
        doc.add_page_break()
        _add_section_heading(doc, f"{entry_name}:", heading_pPr)
        for fb, fname, mime in file_list:
            _add_file(doc, fb, fname, mime)

    # ── 14. Limitations ──────────────────────────────────────────────────────
    doc.add_page_break()
    _add_section_heading(doc, "Limitations:", heading_pPr)
    doc.add_paragraph(
        "This report is based on a visual only inspection conducted at the time of the site visit. "
        "Observations are limited to visible and accessible components only. "
        "No destructive testing or detailed structural analysis was performed."
    )

    # ── 15. Table of Contents — sync cached entries to the real generated
    # headings (see _rewrite_toc_entries docstring for why this is needed
    # in addition to the updateFields flag below).
    _rewrite_toc_entries(doc, _collect_headings_for_toc(doc))

    # Force Word to refresh fields (e.g. the Table of Contents) on open, so
    # newly appended headings show up without the user manually pressing F9.
    _insert_update_fields(doc.settings.element)

    buf = io.BytesIO()
    doc.save(buf)
    return _strip_orphaned_images(buf.getvalue())
