"""
Test: generate a report using the exact data from the original template.
The output should look identical (or very close) to the original.

Run:
  python tests/test_with_template_data.py
Output saved to:
  tests/TEST_OUTPUT_TN20133-A.docx
Open that file and the original template side-by-side in Word to compare.
"""

import sys
import io
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from utils.docx_builder import build_report

# ── Dummy image (grey square — stands in for real photos) ─────────────────────
def _dummy_image() -> bytes:
    try:
        from PIL import Image
        img = Image.new("RGB", (400, 300), color=(160, 160, 160))
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        return buf.getvalue()
    except ImportError:
        # 1×1 white JPEG (minimal valid file)
        return (
            b'\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01\x01\x00\x00\x01\x00\x01\x00\x00'
            b'\xff\xdb\x00C\x00\x08\x06\x06\x07\x06\x05\x08\x07\x07\x07\t\t'
            b'\x08\n\x0c\x14\r\x0c\x0b\x0b\x0c\x19\x12\x13\x0f\x14\x1d\x1a'
            b'\x1f\x1e\x1d\x1a\x1c\x1c $.\' ",#\x1c\x1c(7),01444\x1f\'9=82<.342\x1e'
            b'\x1f\x1f\x1f\x1f\x1f\x1f\x1f\x1f\x1f\xff\xc0\x00\x0b\x08\x00'
            b'\x01\x00\x01\x01\x01\x11\x00\xff\xc4\x00\x1f\x00\x00\x01\x05'
            b'\x01\x01\x01\x01\x01\x01\x00\x00\x00\x00\x00\x00\x00\x00\x01'
            b'\x02\x03\x04\x05\x06\x07\x08\t\n\x0b\xff\xc4\x00\xb5\x10\x00'
            b'\x02\x01\x03\x03\x02\x04\x03\x05\x05\x04\x04\x00\x00\x01}\x01'
            b'\x02\x03\x00\x04\x11\x05\x12!1A\x06\x13Qa\x07"q\x142\x81\x91'
            b'\xa1\x08#B\xb1\xc1\x15R\xd1\xf0$3br\x82\t\n\x16\x17\x18\x19'
            b'\x1a%&\'()*456789:CDEFGHIJSTUVWXYZ'
            b'cdefghijstuvwxyz\x83\x84\x85\x86\x87\x88\x89\x8a\x92\x93\x94'
            b'\x95\x96\x97\x98\x99\x9a\xa2\xa3\xa4\xa5\xa6\xa7\xa8\xa9\xaa'
            b'\xb2\xb3\xb4\xb5\xb6\xb7\xb8\xb9\xba\xc2\xc3\xc4\xc5\xc6\xc7'
            b'\xc8\xc9\xca\xd2\xd3\xd4\xd5\xd6\xd7\xd8\xd9\xda\xe1\xe2\xe3'
            b'\xe4\xe5\xe6\xe7\xe8\xe9\xea\xf1\xf2\xf3\xf4\xf5\xf6\xf7\xf8'
            b'\xf9\xfa\xff\xda\x00\x08\x01\x01\x00\x00?\x00\xfb\xf8\xff\xd9'
        )


DUMMY_IMG   = _dummy_image()
DUMMY_MIME  = "image/jpeg"

# ── Exact data from the original template ─────────────────────────────────────
ORIGINAL_INFO = {
    "report_date":         "June 9, 2026",
    "client":              "SBA Communication Corporation",
    "site_name":           "Hillvale, TN Tower",
    "site_number":         "TN20133-A",
    "carrier_name":        "AT&T",
    "site_address":        "741 Hillvale Road\nAndersonville, TN 37705 (Anderson County)",
    "gps_coords":          "36.1878, -84.0975",
    "tower_type":          "Self Support",
    "tower_height":        "300’",   # curly apostrophe, as in template
    "observation_date":    "08/06/2026",
    "drawing_date":        "01/08/2026",
    "drawing_sheets":      "S-02, S-03",
    "general_contractor":  "Vinco, Inc",
    "gc_contact":          "",
    "project_description": (
        "Post modification structural reinforcement and verification of "
        "tower components as per approved design drawings."
    ),
    "job_no": "TN20133-A",
}

ORIGINAL_MODS = [
    {
        "mod_id":      "M1",
        "description": 'Adding redundant diagonal bracing of L2x2x3/16" (A572-50) Angle at each face.',
        "elevation":   "200.0’-180.0’",
    },
    {
        "mod_id":      "M2",
        "description": 'Adding redundant diagonal bracing of L2x2x3/16" (A572-50) Angle at each face.',
        "elevation":   "180.0’-160.0’",
    },
]

# Field observations pre-filled so we skip the Gemini API call during testing
ORIGINAL_FIELD_OBS = {
    "Structural Member Verification": [
        "Member sizes were verified, including confirmation of installed member dimensions in accordance with the approved drawings.",
        "Installation elevations were checked against the design drawings to confirm proper placement.",
    ],
    "Connection and Installation Verification": [
        "Connection types were verified in accordance with the approved details.",
        "Connection spacing and configuration were checked against design requirements.",
    ],
    "Modification Installation": [
        "Reinforcement members were observed to be installed at the required locations and elevations as per the modification details.",
        "Installation was reviewed for general conformance with the design intent and approved drawings.",
    ],
    "Connection and Welding": [
        "Connections and welded joints were visually inspected for general workmanship, continuity, and alignment.",
        "Weld surfaces were reviewed for visible defects, including cracks or discontinuities.",
    ],
    "Alignment and Eccentricity": [
        "Installed members and connections were reviewed for alignment and fit-up.",
        "No significant misalignment or unintended eccentricity was observed.",
    ],
    "Coating and Protection": [
        "Protective coatings and touch-up applications were reviewed for general coverage and condition.",
        "Coating was observed to be applied where required to maintain corrosion protection.",
    ],
    "Interference Check": [
        "Installed modifications were reviewed for potential interference with existing tower members or components.",
        "No significant interference was observed.",
    ],
    "Final Verification": [
        "The modification work was reviewed for overall completeness and general conformance with the approved drawings.",
        "Photographic documentation was obtained during the inspection.",
    ],
}

# Photos: 4 per leg to demo 2×2 grid layout; 2 for some legs to show adaptive behaviour
PHOTOS = {
    ("M1", "Leg A"): [
        (DUMMY_IMG, "M1 Leg A — Photo 1: installed diagonal bracing at 200’-180’, overview.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg A — Photo 2: close-up of bolt connection at top node.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg A — Photo 3: weld joint inspection, no visible defects.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg A — Photo 4: elevation measurement verification.", DUMMY_MIME),
    ],
    ("M1", "Leg B"): [
        (DUMMY_IMG, "M1 Leg B — Photo 1: diagonal L-angle bracing installed.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg B — Photo 2: bolt pattern and connection detail.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg B — Photo 3: coating applied to installed members.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg B — Photo 4: general view of completed Leg B work.", DUMMY_MIME),
    ],
    ("M1", "Leg C"): [
        (DUMMY_IMG, "M1 Leg C — Photo 1: redundant bracing installed per design.", DUMMY_MIME),
        (DUMMY_IMG, "M1 Leg C — Photo 2: connection alignment verified.", DUMMY_MIME),
    ],
    ("M2", "Leg A"): [
        (DUMMY_IMG, "M2 Leg A — Photo 1: diagonal bracing at 180’-160’, front view.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg A — Photo 2: fastener count confirmed per drawings.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg A — Photo 3: weld inspection, no discontinuities.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg A — Photo 4: cold galvanization touch-up visible.", DUMMY_MIME),
    ],
    ("M2", "Leg B"): [
        (DUMMY_IMG, "M2 Leg B — Photo 1: completed installation overview.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg B — Photo 2: weld connections inspected.", DUMMY_MIME),
    ],
    ("M2", "Leg C"): [
        (DUMMY_IMG, "M2 Leg C — Photo 1: installed members consistent with approved drawings.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg C — Photo 2: no eccentricity observed.", DUMMY_MIME),
        (DUMMY_IMG, "M2 Leg C — Photo 3: coating coverage confirmed.", DUMMY_MIME),
    ],
}

SPECIAL_PHOTOS = {
    "Overall Tower View": [
        (DUMMY_IMG, "Overall tower view — south face, completed structural modifications visible.", DUMMY_MIME),
        (DUMMY_IMG, "Overall tower view — north face, all legs inspected.", DUMMY_MIME),
    ],
    "Field Measurements": [
        (DUMMY_IMG, "Field measurement — diagonal member width confirmed 2x2x3/16\".", DUMMY_MIME),
        (DUMMY_IMG, "Field measurement — elevation mark at 180’ verified with tape.", DUMMY_MIME),
    ],
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _check(doc_bytes: bytes) -> list[str]:
    """Open the generated report and verify key text fields are present."""
    doc = Document(io.BytesIO(doc_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    checks = [
        ("Report date",           "June 9, 2026"),
        ("Client name",           "SBA Communication Corporation"),
        ("Site name",             "Hillvale, TN Tower"),
        ("Site number",           "TN20133-A"),
        ("Carrier",               "AT&T"),
        ("Site address",          "741 Hillvale Road"),
        ("GPS",                   "36.1878"),
        ("Tower height",          "300"),
        ("Observation date",      "08/06/2026"),
        ("Drawing sheets",        "S-02, S-03"),
        ("Drawing date",          "01/08/2026"),
        ("General contractor",    "Vinco, Inc"),
        ("M1 in table",           "M1"),
        ("M2 in table",           "M2"),
        ("Field obs heading",     "Field Observation Details"),
        ("Structural Member Ver.", "Structural Member Verification"),
        ("No deficiencies",       "No visible deficiencies"),
        ("Photo section",         "On-Site Inspection Photographs"),
        ("Limitations",           "Limitations"),
    ]

    failures = []
    for label, expected in checks:
        if expected not in all_text:
            failures.append(f"MISSING — {label}: '{expected}'")
    return failures


# ── Main ──────────────────────────────────────────────────────────────────────

def run():
    print("Building report with original template data...")
    print()

    data = {
        "info":              ORIGINAL_INFO,
        "modifications":     ORIGINAL_MODS,
        "photos":            PHOTOS,
        "special_photos":    SPECIAL_PHOTOS,
        "documents":         {},          # no docs — [Not provided] placeholders appear
        "certificates":      [],
        "deficiencies":      "",
        "no_deficiencies":   True,
        "tower_type":        "Self Support",
        "num_guys":          3,
        "field_observations": ORIGINAL_FIELD_OBS,
    }

    try:
        report_bytes = build_report(data)
    except Exception as e:
        print(f"[FAIL] build_report() raised an exception:\n  {e}")
        import traceback
        traceback.print_exc()
        return

    out_path = Path(__file__).parent / "TEST_OUTPUT_TN20133-A_v3.docx"
    out_path.write_bytes(report_bytes)
    print(f"Report saved to: {out_path}")
    print(f"File size: {len(report_bytes):,} bytes")
    print()

    # Verify
    failures = _check(report_bytes)

    if failures:
        print(f"[FAIL] {len(failures)} check(s) failed:")
        for f in failures:
            print(f"  FAIL: {f}")
    else:
        print("[PASS] All field checks passed.")

    print()
    print("Open the file in Word and compare to the original template.")


if __name__ == "__main__":
    run()
