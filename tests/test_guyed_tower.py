"""
Test: generate a full Guyed Tower PMI closeout report to verify the
guyed-specific pipeline end-to-end (positions, docs, special photos,
field observation sections) and the photo grid layout.

Run:
  python tests/test_guyed_tower.py
Output saved to:
  tests/TEST_OUTPUT_GUYED.docx
"""

import sys
import io
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from utils.docx_builder import build_report, OBS_SECTIONS_BY_TOWER
from utils.validators import get_required_docs, get_special_photos, check_missing_items


def _dummy_image(color=(160, 160, 160)) -> bytes:
    from PIL import Image
    img = Image.new("RGB", (400, 300), color=color)
    buf = io.BytesIO()
    img.save(buf, format="JPEG")
    return buf.getvalue()


DUMMY_IMG  = _dummy_image()
DUMMY_MIME = "image/jpeg"

INFO = {
    "report_date":         "July 2, 2026",
    "client":              "SBA Communication Corporation",
    "site_name":           "Blount County Guyed Tower",
    "site_number":         "SCO1980-A",
    "carrier_name":        "Verizon",
    "site_address":        "1980 Sevierville Road\nMaryville, TN 37801 (Blount County)",
    "gps_coords":          "35.7565, -83.9705",
    "tower_type":          "Guyed",
    "tower_height":        "400'",
    "observation_date":    "05/22/2026",
    "drawing_date":        "04/15/2026",
    "drawing_sheets":      "S-01, S-02, S-03",
    "general_contractor":  "Vinco, Inc",
    "gc_contact":          "",
    "project_description": (
        "Post modification structural reinforcement and verification of "
        "guyed tower components as per approved design drawings."
    ),
    "job_no": "SCO1980-A",
}

MODS = [
    {
        "mod_id":      "M1",
        "description": 'Adding redundant guy wire anchor reinforcement plate at Level 3 anchor.',
        "elevation":   "250.0'-230.0'",
    },
    {
        "mod_id":      "M2",
        "description": 'Replacing corroded guy wire hardware and tensioning to design specification.',
        "elevation":   "150.0'-130.0'",
    },
]

NUM_GUYS = 3  # Guy 1, Guy 2, Guy 3

FIELD_OBS = {
    section: [
        f"Sample observation 1 for {section.lower()} on the guyed tower.",
        f"Sample observation 2 for {section.lower()}, no deficiencies noted.",
    ]
    for section in OBS_SECTIONS_BY_TOWER["Guyed"]
}

# Photos: mix of 4 (full 2x2 grid), 3 (2+1), and 2 (one row) to exercise the grid logic
PHOTOS = {}
photo_counts = {("M1", "Guy 1"): 4, ("M1", "Guy 2"): 4, ("M1", "Guy 3"): 3,
                 ("M2", "Guy 1"): 4, ("M2", "Guy 2"): 2, ("M2", "Guy 3"): 4}
for (mid, guy), count in photo_counts.items():
    PHOTOS[(mid, guy)] = [
        (DUMMY_IMG, f"{mid} {guy} — Photo {i+1}: guy wire anchor inspection detail.", DUMMY_MIME)
        for i in range(count)
    ]

SPECIAL_PHOTOS = {
    "Tension Gauge Photos": [
        (DUMMY_IMG, "Tension gauge reading — Guy 1, anchor point, 8,200 lbs.", DUMMY_MIME),
        (DUMMY_IMG, "Tension gauge reading — Guy 2, anchor point, 8,150 lbs.", DUMMY_MIME),
        (DUMMY_IMG, "Tension gauge reading — Guy 3, anchor point, 8,300 lbs.", DUMMY_MIME),
        (DUMMY_IMG, "Tension gauge calibration certificate photo.", DUMMY_MIME),
        (DUMMY_IMG, "Tension gauge — overall setup view.", DUMMY_MIME),
    ],
    "Overall Tower View": [
        (DUMMY_IMG, "Overall tower view — north face, guy anchors visible.", DUMMY_MIME),
        (DUMMY_IMG, "Overall tower view — south face.", DUMMY_MIME),
    ],
}

# Guyed-specific docs — dummy "pages" so every section renders content instead of [Not provided]
DOCUMENTS = {
    "as_built_drawings":     [(DUMMY_IMG, "as_built_p1.jpg", DUMMY_MIME)],
    "material_cert":         [(DUMMY_IMG, "material_cert_p1.jpg", DUMMY_MIME)],
    "packing_slips":         [(DUMMY_IMG, "packing_slip_p1.jpg", DUMMY_MIME)],
    "tension_report":        [(DUMMY_IMG, "tension_report_p1.jpg", DUMMY_MIME)],
    "plumb_twist_report":    [(DUMMY_IMG, "plumb_twist_p1.jpg", DUMMY_MIME)],
    "fabrication_submittal": [(DUMMY_IMG, "fab_submittal_p1.jpg", DUMMY_MIME)],
    "cold_galv_letter":      [(DUMMY_IMG, "cold_galv_p1.jpg", DUMMY_MIME)],
}

EXTRA_DOCUMENTS = [
    ("Welder Certification", [(DUMMY_IMG, "welder_cert.jpg", DUMMY_MIME)]),
    ("Tension Gauge Calibration Certificate", [(DUMMY_IMG, "tension_gauge_cal.jpg", DUMMY_MIME)]),
]


def run():
    print("=== Guyed Tower pipeline sanity checks ===")
    assert "fabrication_letter" not in [k for k, _ in get_required_docs("Guyed")], \
        "Guyed docs should NOT include Fabrication Letter"
    assert ("packing_slips", "Packing Slips") in get_required_docs("Guyed")
    assert ("tension_report", "Tension Report") in get_required_docs("Guyed")
    assert ("plumb_twist_report", "Plumb & Twist Report") in get_required_docs("Guyed")
    assert "Tension Gauge Photos" in get_special_photos("Guyed")
    assert "Guy Wire Verification" in OBS_SECTIONS_BY_TOWER["Guyed"]
    assert "Connection and Welding" not in OBS_SECTIONS_BY_TOWER["Guyed"]
    assert "Coating and Protection" not in OBS_SECTIONS_BY_TOWER["Guyed"]
    print("[PASS] validators/docx_builder guyed-specific config looks correct")

    check_data = {
        "tower_type":    "Guyed",
        "modifications": MODS,
        "num_guys":      NUM_GUYS,
        "photos":        dict(PHOTOS),
        "documents":     DOCUMENTS,
        "extra_documents": [{"name": n, "files": f} for n, f in EXTRA_DOCUMENTS],
    }
    for label, pl in SPECIAL_PHOTOS.items():
        check_data["photos"][("special", label)] = pl
    missing = check_missing_items(check_data)
    total_missing = sum(len(v) for v in missing.values())
    print(f"Missing items with full data supplied: {missing}")
    assert total_missing == 0, f"Expected zero missing items, got: {missing}"
    print("[PASS] check_missing_items reports nothing missing when all data supplied")

    print()
    print("Building guyed tower report...")
    data = {
        "info":              INFO,
        "modifications":     MODS,
        "photos":            PHOTOS,
        "special_photos":    SPECIAL_PHOTOS,
        "documents":         DOCUMENTS,
        "extra_documents":   EXTRA_DOCUMENTS,
        "deficiencies":      "",
        "no_deficiencies":   True,
        "tower_type":        "Guyed",
        "num_guys":          NUM_GUYS,
        "field_observations": FIELD_OBS,
    }

    try:
        report_bytes = build_report(data)
    except Exception as e:
        print(f"[FAIL] build_report() raised an exception:\n  {e}")
        import traceback
        traceback.print_exc()
        return

    out_path = Path(__file__).parent / "TEST_OUTPUT_GUYED.docx"
    out_path.write_bytes(report_bytes)
    print(f"Report saved to: {out_path}")
    print(f"File size: {len(report_bytes):,} bytes")
    print()

    doc = Document(io.BytesIO(report_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    checks = [
        ("Report date",            "July 2, 2026"),
        ("Site number",            "SCO1980-A"),
        ("Tower height",           "400"),
        ("Tower type label",       "Guyed Tower"),
        ("M1 in table",            "M1"),
        ("M2 in table",            "M2"),
        ("Guy 1 photo row",        "Guy 1"),
        ("Guy 3 photo row",        "Guy 3"),
        ("Guy Wire Verification",  "Guy Wire Verification"),
        ("No Connection/Welding heading (guyed shouldn't have it)", None),
        ("Tension Gauge special photo heading", "Tension Gauge Photographs"),
        ("Overall Tower View heading",          "Overall Tower Views"),
        ("Packing Slips doc heading",           "Packing Slips:"),
        ("Tension Report doc heading",          "Tension Report:"),
        ("Plumb & Twist Report doc heading",    "Plumb & Twist Report:"),
        ("Field obs heading",       "Field Observation Details"),
        ("Custom cert heading 1",   "Welder Certification:"),
        ("Custom cert heading 2",   "Tension Gauge Calibration Certificate:"),
        ("No generic Certificates heading", None),
        ("Limitations",             "Limitations"),
    ]

    failures = []
    for label, expected in checks:
        if expected is None:
            continue
        if expected not in all_text:
            failures.append(f"MISSING — {label}: '{expected}'")

    if "Connection and Welding:" in all_text or "Connection/Welding" in all_text:
        failures.append("Guyed report unexpectedly contains a Connection/Welding field-obs heading")
    if "Certificates :" in all_text or "Certificates:" in all_text:
        failures.append("Report unexpectedly still contains the old generic 'Certificates' heading")
    if "Coating and Protection:" in all_text or "Coating/Protection" in all_text:
        failures.append("Guyed report unexpectedly contains a Coating/Protection field-obs heading")
    if "Fabrication Letter:" in all_text:
        failures.append("Guyed report unexpectedly contains a Fabrication Letter doc section")

    if failures:
        print(f"[FAIL] {len(failures)} check(s) failed:")
        for f in failures:
            print(f"  FAIL: {f}")
    else:
        print("[PASS] All guyed-tower content checks passed.")


if __name__ == "__main__":
    run()
