"""Reads the original template and prints its content for test data extraction."""
import sys
sys.path.insert(0, r"C:\Users\VBbbbbbvvv\Documents\Claude Agent Project")

from docx import Document
from pathlib import Path

TEMPLATE = Path(r"C:\Users\VBbbbbbvvv\Documents\Claude Agent Project\TN20133-A PMI Closeout Report .docx")
doc = Document(TEMPLATE)

print("=== BODY PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i:3d}] {repr(p.text[:150])}")

print()
print("=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ---")
    for ri, row in enumerate(table.rows):
        cells = [c.text[:80] for c in row.cells]
        print(f"  Row {ri}: {cells}")
